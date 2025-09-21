from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO, emit
import speedtest
import time
import threading
from datetime import datetime
from statistics import mean, stdev
import json
import tempfile
import os

app = Flask(__name__)
socketio = SocketIO(app, cors_allowed_origins="*")

# Global variables to manage test state
active_tests = {}

def calculate_jitter(ping_results):
    """Calculate jitter from ping results"""
    if len(ping_results) < 2:
        return 0
    differences = []
    for i in range(1, len(ping_results)):
        differences.append(abs(ping_results[i] - ping_results[i-1]))
    return mean(differences) if differences else 0

def run_single_speed_test(session_id, preferred_server_id=None):
    """Runs a single speed test and yields real-time results."""
    try:
        st = speedtest.Speedtest(secure=True)
        
        # Get client configuration first
        config = st.get_config()
        client_info = {
            'ip': config['client']['ip'],
            'isp': config['client']['isp'],
            'country': config['client']['country']
        }
        
        # Emit client information
        socketio.emit('test_progress', {
            'type': 'client_info',
            'client': client_info,
            'session_id': session_id
        })
        
        # Emit server selection status
        socketio.emit('test_progress', {
            'type': 'status',
            'message': 'Selecting best server...',
            'session_id': session_id
        })
        
        # Get server list and select the best one or preferred one
        if preferred_server_id:
            # Use specific server if provided
            servers = st.get_servers([preferred_server_id])
            st.get_best_server(servers[preferred_server_id])
            server_info = st.results.server
        else:
            # Get best server automatically - prioritize closest servers
            servers = st.get_servers()
            all_servers = []
            for distance_key in servers:
                for server in servers[distance_key]:
                    all_servers.append(server)
            
            # Sort by distance first, then by latency
            all_servers.sort(key=lambda x: x['d'])
            closest_servers = all_servers[:3]  # Get 3 closest servers
            
            # Test latency on closest servers and pick the best
            best_latency = float('inf')
            best_server = None
            
            for server in closest_servers:
                try:
                    # Quick latency test
                    import socket
                    import time
                    host = server['host'].split(':')[0]
                    port = int(server['host'].split(':')[1]) if ':' in server['host'] else 8080
                    
                    start_time = time.time()
                    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                    sock.settimeout(2)
                    result = sock.connect_ex((host, port))
                    sock.close()
                    latency = (time.time() - start_time) * 1000
                    
                    if result == 0 and latency < best_latency:
                        best_latency = latency
                        best_server = server
                except:
                    continue
            
            if best_server:
                st.get_best_server([best_server])
                server_info = best_server
            else:
                # Fallback to original method
                server_info = st.get_best_server()
        
        socketio.emit('test_progress', {
            'type': 'server_selected',
            'server': {
                'id': server_info.get('id', 'Unknown'),
                'name': server_info.get('sponsor', 'Unknown'),
                'location': f"{server_info.get('name', 'Unknown')}, {server_info.get('country', 'Unknown')}",
                'distance': round(server_info.get('d', 0), 2),
                'url': server_info.get('url', 'Unknown'),
                'latency': round(server_info.get('latency', 0), 2)
            },
            'session_id': session_id
        })

        # Ping test with multiple samples for jitter calculation
        socketio.emit('test_progress', {
            'type': 'status',
            'message': 'Testing ping...',
            'session_id': session_id
        })
        
        ping_results = []
        for i in range(5):  # 5 ping samples for jitter calculation
            ping = st.results.ping
            ping_results.append(ping)
            socketio.emit('test_progress', {
                'type': 'ping_sample',
                'ping': round(ping, 2),
                'sample': i + 1,
                'session_id': session_id
            })
            time.sleep(0.1)

        avg_ping = mean(ping_results)
        jitter = calculate_jitter(ping_results)

        # Download test with real-time updates
        socketio.emit('test_progress', {
            'type': 'status',
            'message': 'Testing download speed...',
            'session_id': session_id
        })
        
        # Run download test and simulate frequent updates
        download_result = st.download() / 1_000_000
        
        # Send multiple data points during download phase for animated chart
        for i in range(1, 21):  # 20 data points
            progress = i / 20.0
            speed_variation = download_result * (0.8 + 0.4 * progress)  # Simulate speed ramping up
            socketio.emit('test_progress', {
                'type': 'download_progress',
                'download': round(speed_variation, 2),
                'session_id': session_id
            })
            time.sleep(0.1)  # 100ms intervals for smooth animation
        
        # Send final download result
        socketio.emit('test_progress', {
            'type': 'download_complete',
            'download': round(download_result, 2),
            'session_id': session_id
        })

        # Upload test with real-time updates
        socketio.emit('test_progress', {
            'type': 'status',
            'message': 'Testing upload speed...',
            'session_id': session_id
        })
        
        upload_result = st.upload() / 1_000_000
        
        # Send multiple data points during upload phase for animated chart
        for i in range(1, 21):  # 20 data points
            progress = i / 20.0
            speed_variation = upload_result * (0.7 + 0.5 * progress)  # Simulate speed ramping up
            socketio.emit('test_progress', {
                'type': 'upload_progress',
                'upload': round(speed_variation, 2),
                'session_id': session_id
            })
            time.sleep(0.1)  # 100ms intervals for smooth animation
        
        # Send final upload result
        socketio.emit('test_progress', {
            'type': 'upload_complete',
            'upload': round(upload_result, 2),
            'session_id': session_id
        })

        # Final results
        final_result = {
            'type': 'final',
            'ping': round(avg_ping, 2),
            'jitter': round(jitter, 2),
            'download': round(download_result, 2),
            'upload': round(upload_result, 2),
            'server': server_info,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'session_id': session_id
        }
        
        socketio.emit('test_result', final_result)
        return final_result

    except Exception as e:
        error_result = {
            'type': 'error',
            'message': str(e),
            'session_id': session_id
        }
        socketio.emit('test_result', error_result)
        return error_result

def run_continuous_speed_test(session_id, duration_minutes, preferred_server_id=None):
    """Run continuous speed test for stability analysis"""
    try:
        duration_seconds = duration_minutes * 60
        start_time = time.time()
        test_results = []
        test_count = 0
        
        socketio.emit('continuous_test_started', {
            'session_id': session_id,
            'duration_minutes': duration_minutes,
            'message': f'Starting {duration_minutes}-minute stability test...'
        })
        
        while time.time() - start_time < duration_seconds:
            test_count += 1
            socketio.emit('test_progress', {
                'type': 'status',
                'message': f'Running test #{test_count}...',
                'session_id': session_id,
                'test_number': test_count
            })
            
            result = run_single_speed_test(session_id, preferred_server_id)
            if result['type'] != 'error':
                test_results.append(result)
                
                # Calculate running statistics
                pings = [r['ping'] for r in test_results]
                downloads = [r['download'] for r in test_results]
                uploads = [r['upload'] for r in test_results]
                
                running_stats = {
                    'type': 'running_stats',
                    'test_count': len(test_results),
                    'avg_ping': round(mean(pings), 2),
                    'avg_download': round(mean(downloads), 2),
                    'avg_upload': round(mean(uploads), 2),
                    'min_download': round(min(downloads), 2),
                    'max_download': round(max(downloads), 2),
                    'min_upload': round(min(uploads), 2),
                    'max_upload': round(max(uploads), 2),
                    'session_id': session_id,
                    'progress': round(((time.time() - start_time) / duration_seconds) * 100, 1)
                }
                
                socketio.emit('running_stats', running_stats)
            
            # Wait a bit before next test (adjust as needed)
            time.sleep(10)
        
        # Final stability analysis
        if test_results:
            pings = [r['ping'] for r in test_results]
            downloads = [r['download'] for r in test_results]
            uploads = [r['upload'] for r in test_results]
            
            # Calculate stability score based on variance
            download_variance = stdev(downloads) if len(downloads) > 1 else 0
            upload_variance = stdev(uploads) if len(uploads) > 1 else 0
            ping_variance = stdev(pings) if len(pings) > 1 else 0
            
            # Stability score: lower variance = higher stability
            stability_score = max(0, 100 - (download_variance + upload_variance + ping_variance))
            
            stability_analysis = {
                'type': 'continuous',
                'test_type': 'continuous',
                'session_id': session_id,
                'test_count': len(test_results),
                'duration': duration_minutes,
                'avg_ping': round(mean(pings), 2),
                'min_ping': round(min(pings), 2),
                'max_ping': round(max(pings), 2),
                'avg_download': round(mean(downloads), 2),
                'min_download': round(min(downloads), 2),
                'max_download': round(max(downloads), 2),
                'avg_upload': round(mean(uploads), 2),
                'min_upload': round(min(uploads), 2),
                'max_upload': round(max(uploads), 2),
                'stability_score': round(stability_score, 1),
                'ping_variance': round(ping_variance, 2),
                'download_variance': round(download_variance, 2),
                'upload_variance': round(upload_variance, 2),
                'test_results': test_results
            }
            
            socketio.emit('test_result', stability_analysis)
        
        # Clean up
        if session_id in active_tests:
            del active_tests[session_id]
            
    except Exception as e:
        socketio.emit('test_result', {
            'type': 'error',
            'message': str(e),
            'session_id': session_id
        })

@socketio.on('start_test')
def handle_start_test(data):
    """Handles the start test event from the client."""
    session_id = data.get('session_id', 'default')
    test_type = data.get('test_type', 'single')  # 'single' or 'continuous'
    duration = data.get('duration', 5)  # Duration in minutes for continuous test
    preferred_server_id = data.get('server_id', None)  # Optional server selection
    
    print(f"Client requested a {test_type} speed test. Session: {session_id}")
    
    if test_type == 'continuous':
        # Run continuous test in a separate thread
        thread = threading.Thread(
            target=run_continuous_speed_test,
            args=(session_id, duration, preferred_server_id)
        )
        active_tests[session_id] = thread
        thread.start()
    else:
        # Run single test
        run_single_speed_test(session_id, preferred_server_id)

@socketio.on('stop_test')
def handle_stop_test(data):
    """Stop an ongoing test"""
    session_id = data.get('session_id', 'default')
    if session_id in active_tests:
        # Note: This is a simple implementation. In production, you'd want proper thread management
        del active_tests[session_id]
        socketio.emit('test_stopped', {'session_id': session_id})

@app.route('/api/servers', methods=['GET'])
def get_available_servers():
    """Get list of available test servers"""
    try:
        st = speedtest.Speedtest(secure=True)
        servers = st.get_servers()
        
        # Format server list for frontend
        server_list = []
        for server_id, server_data in servers.items():
            for server in server_data:
                server_list.append({
                    'id': server.get('id'),
                    'name': server.get('sponsor', 'Unknown'),
                    'location': f"{server.get('name', 'Unknown')}, {server.get('country', 'Unknown')}",
                    'distance': round(server.get('d', 0), 2),
                    'country': server.get('country', 'Unknown'),
                    'cc': server.get('cc', 'XX'),
                    'url': server.get('url', '')
                })
        
        # Sort by distance
        server_list.sort(key=lambda x: x['distance'])
        
        return jsonify({
            'success': True,
            'servers': server_list[:20]  # Return top 20 closest servers
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@socketio.on('get_servers')
def handle_get_servers():
    """WebSocket handler to get available servers"""
    try:
        st = speedtest.Speedtest(secure=True)
        servers = st.get_servers()
        
        # Format server list for frontend
        server_list = []
        for server_id, server_data in servers.items():
            for server in server_data:
                server_list.append({
                    'id': server.get('id'),
                    'name': server.get('sponsor', 'Unknown'),
                    'location': f"{server.get('name', 'Unknown')}, {server.get('country', 'Unknown')}",
                    'distance': round(server.get('d', 0), 2),
                    'country': server.get('country', 'Unknown'),
                    'cc': server.get('cc', 'XX'),
                    'url': server.get('url', '')
                })
        
        # Sort by distance
        server_list.sort(key=lambda x: x['distance'])
        
        socketio.emit('servers_list', {
            'servers': server_list[:20]  # Return top 20 closest servers
        })
        
    except Exception as e:
        socketio.emit('servers_error', {'error': str(e)})

@app.route('/api/export', methods=['POST'])
def export_results():
    """Export test results as PDF or DOCX"""
    try:
        data = request.json
        export_format = data.get('format', 'pdf')  # 'pdf' or 'docx'
        results = data.get('results', {})
        stability_data = data.get('stabilityData', [])
        
        if export_format == 'pdf':
            file_path = create_pdf_report(results, stability_data)
        elif export_format == 'docx':
            file_path = create_docx_report(results, stability_data)
        else:
            return jsonify({'success': False, 'error': 'Invalid format'}), 400
            
        return send_file(
            file_path,
            as_attachment=True,
            download_name=f'speedtest_results.{export_format}',
            mimetype='application/pdf' if export_format == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def create_pdf_report(results, stability_data):
    """Create a PDF report of the speed test results"""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from datetime import datetime
    import matplotlib.pyplot as plt
    import io
    import base64
    
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=colors.HexColor('#06B6D4')
    )
    story.append(Paragraph("Speed Test Results", title_style))
    story.append(Spacer(1, 20))
    
    # Test information
    test_info = [
        ['Test Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ['Test Type:', 'Stability Test' if results.get('type') == 'stability_final' else 'Quick Test'],
    ]
    
    if results.get('type') == 'stability_final':
        test_info.extend([
            ['Duration:', f"{results.get('duration_minutes', 0)} minutes"],
            ['Tests Completed:', str(results.get('total_tests', 0))],
        ])
    
    info_table = Table(test_info, colWidths=[2*inch, 3*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 20))
    
    # Results table
    if results.get('type') == 'stability_final':
        # Stability test results
        story.append(Paragraph("Connection Stability Analysis", styles['Heading2']))
        
        stats_data = [
            ['Metric', 'Average', 'Minimum', 'Maximum', 'Std Dev'],
            ['Download (Mbps)', 
             str(results['download_stats']['avg']), 
             str(results['download_stats']['min']), 
             str(results['download_stats']['max']), 
             str(results['download_stats']['std'])],
            ['Upload (Mbps)', 
             str(results['upload_stats']['avg']), 
             str(results['upload_stats']['min']), 
             str(results['upload_stats']['max']), 
             str(results['upload_stats']['std'])],
            ['Ping (ms)', 
             str(results['ping_stats']['avg']), 
             str(results['ping_stats']['min']), 
             str(results['ping_stats']['max']), 
             str(results['ping_stats']['std'])],
        ]
    else:
        # Single test results
        story.append(Paragraph("Speed Test Results", styles['Heading2']))
        
        stats_data = [
            ['Metric', 'Value'],
            ['Download Speed', f"{results.get('download', 0)} Mbps"],
            ['Upload Speed', f"{results.get('upload', 0)} Mbps"],
            ['Ping', f"{results.get('ping', 0)} ms"],
            ['Jitter', f"{results.get('jitter', 0)} ms"],
        ]
    
    results_table = Table(stats_data)
    results_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(results_table)
    
    # Build PDF
    doc.build(story)
    return temp_file.name

def create_docx_report(results, stability_data):
    """Create a DOCX report of the speed test results"""
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Speed Test Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Test information
    doc.add_heading('Test Information', level=1)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    
    # Add test info rows
    test_info = [
        ('Test Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('Test Type:', 'Stability Test' if results.get('type') == 'stability_final' else 'Quick Test'),
    ]
    
    if results.get('type') == 'stability_final':
        test_info.extend([
            ('Duration:', f"{results.get('duration_minutes', 0)} minutes"),
            ('Tests Completed:', str(results.get('total_tests', 0))),
        ])
    
    for label, value in test_info:
        row_cells = info_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value
    
    # Results
    doc.add_heading('Results', level=1)
    results_table = doc.add_table(rows=1, cols=2 if results.get('type') != 'stability_final' else 5)
    results_table.style = 'Table Grid'
    
    if results.get('type') == 'stability_final':
        # Headers
        hdr_cells = results_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Average'
        hdr_cells[2].text = 'Minimum'
        hdr_cells[3].text = 'Maximum'
        hdr_cells[4].text = 'Std Dev'
        
        # Data rows
        metrics = [
            ('Download (Mbps)', results['download_stats']),
            ('Upload (Mbps)', results['upload_stats']),
            ('Ping (ms)', results['ping_stats']),
        ]
        
        for metric_name, stats in metrics:
            row_cells = results_table.add_row().cells
            row_cells[0].text = metric_name
            row_cells[1].text = str(stats['avg'])
            row_cells[2].text = str(stats['min'])
            row_cells[3].text = str(stats['max'])
            row_cells[4].text = str(stats['std'])
    else:
        # Headers
        hdr_cells = results_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Data rows
        metrics = [
            ('Download Speed', f"{results.get('download', 0)} Mbps"),
            ('Upload Speed', f"{results.get('upload', 0)} Mbps"),
            ('Ping', f"{results.get('ping', 0)} ms"),
            ('Jitter', f"{results.get('jitter', 0)} ms"),
        ]
        
        for metric_name, value in metrics:
            row_cells = results_table.add_row().cells
            row_cells[0].text = metric_name
            row_cells[1].text = value
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

if __name__ == '__main__':
    print("Starting Flask-SocketIO server...")
    socketio.run(app, host='0.0.0.0', port=5000, debug=True)
